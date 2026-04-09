import sys
sys.path.insert(0, '.')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 表结构数据
SCHEMA = {
    "users": {
        "cn": "用户表",
        "desc": "存储系统用户信息，包括管理员和普通用户。",
        "columns": [
            ("id", "UUID", "否", "uuid_generate_v4()", "主键", "用户唯一标识"),
            ("created_at", "TIMESTAMPTZ", "否", "now()", "", "创建时间"),
            ("updated_at", "TIMESTAMPTZ", "否", "now()", "", "更新时间"),
            ("deleted_at", "TIMESTAMPTZ", "是", "NULL", "", "软删除时间"),
            ("username", "VARCHAR(255)", "否", "NULL", "唯一", "用户名"),
            ("password", "VARCHAR(255)", "是", "NULL", "", "加密密码"),
            ("cellphone", "VARCHAR(45)", "是", "NULL", "唯一", "手机号"),
            ("state", "user_state_type", "否", "enabled", "", "用户状态（enabled/disabled）"),
            ("avatar", "VARCHAR", "否", "NULL", "", "头像路径"),
            ("is_admin", "BOOLEAN", "否", "NULL", "", "是否为管理员"),
        ]
    },
    "categories": {
        "cn": "商品分类表",
        "desc": "存储农产品的分类信息，如蔬菜、水果等。",
        "columns": [
            ("id", "INTEGER", "否", "自增", "主键", "分类唯一标识"),
            ("name", "VARCHAR(50)", "否", "NULL", "", "分类名称"),
            ("created_at", "TIMESTAMPTZ", "是", "now()", "", "创建时间"),
        ]
    },
    "products": {
        "cn": "农产品表",
        "desc": "存储农产品基本信息，关联分类表。",
        "columns": [
            ("id", "INTEGER", "否", "自增", "主键", "产品唯一标识"),
            ("name", "VARCHAR(100)", "否", "NULL", "", "产品名称"),
            ("category_id", "INTEGER", "是", "NULL", "外键→categories.id", "所属分类"),
            ("unit", "VARCHAR(20)", "是", "斤", "", "计量单位"),
            ("remark", "TEXT", "是", "NULL", "", "备注"),
            ("created_at", "TIMESTAMPTZ", "是", "now()", "", "创建时间"),
        ]
    },
    "markets": {
        "cn": "市场表",
        "desc": "存储农产品批发市场信息。",
        "columns": [
            ("id", "INTEGER", "否", "自增", "主键", "市场唯一标识"),
            ("name", "VARCHAR(100)", "否", "NULL", "", "市场名称"),
            ("province", "VARCHAR(50)", "是", "NULL", "", "所在省份"),
            ("city", "VARCHAR(50)", "是", "NULL", "", "所在城市"),
            ("created_at", "TIMESTAMPTZ", "是", "now()", "", "创建时间"),
        ]
    },
    "data_sources": {
        "cn": "数据来源表",
        "desc": "存储爬虫数据来源配置信息。",
        "columns": [
            ("id", "INTEGER", "否", "自增", "主键", "来源唯一标识"),
            ("name", "VARCHAR(50)", "否", "NULL", "", "来源名称"),
            ("url", "TEXT", "是", "NULL", "", "来源URL"),
            ("crawler", "VARCHAR(50)", "是", "NULL", "", "爬虫类名"),
            ("is_active", "BOOLEAN", "是", "true", "", "是否启用"),
            ("created_at", "TIMESTAMPTZ", "是", "now()", "", "创建时间"),
        ]
    },
    "price_records": {
        "cn": "价格记录表",
        "desc": "存储农产品历史价格数据，为系统核心数据表。",
        "columns": [
            ("time", "TIMESTAMPTZ", "否", "NULL", "", "价格记录时间"),
            ("product_id", "INTEGER", "否", "NULL", "外键→products.id", "关联产品"),
            ("market_id", "INTEGER", "是", "NULL", "外键→markets.id", "关联市场"),
            ("source_id", "INTEGER", "是", "NULL", "外键→data_sources.id", "关联数据来源"),
            ("price", "NUMERIC", "是", "NULL", "", "价格"),
            ("min_price", "NUMERIC", "是", "NULL", "", "最低价"),
            ("max_price", "NUMERIC", "是", "NULL", "", "最高价"),
            ("avg_price", "NUMERIC", "是", "NULL", "", "平均价"),
            ("spec", "VARCHAR(100)", "是", "NULL", "", "规格"),
            ("unit", "VARCHAR(20)", "是", "NULL", "", "单位"),
        ]
    },
    "crawl_logs": {
        "cn": "爬虫日志表",
        "desc": "记录每次爬虫任务的执行情况。",
        "columns": [
            ("id", "INTEGER", "否", "自增", "主键", "日志唯一标识"),
            ("source_id", "INTEGER", "是", "NULL", "外键→data_sources.id", "关联数据来源"),
            ("started_at", "TIMESTAMPTZ", "是", "NULL", "", "开始时间"),
            ("finished_at", "TIMESTAMPTZ", "是", "NULL", "", "结束时间"),
            ("status", "VARCHAR(20)", "是", "NULL", "", "执行状态"),
            ("saved_count", "INTEGER", "是", "NULL", "", "入库数量"),
            ("skipped_count", "INTEGER", "是", "NULL", "", "跳过数量"),
            ("error_msg", "TEXT", "是", "NULL", "", "错误信息"),
            ("created_at", "TIMESTAMPTZ", "是", "now()", "", "创建时间"),
            ("updated_at", "TIMESTAMPTZ", "是", "now()", "", "更新时间"),
        ]
    },
    "price_alerts": {
        "cn": "价格预警规则表",
        "desc": "存储用户设置的价格预警规则。",
        "columns": [
            ("id", "INTEGER", "否", "自增", "主键", "规则唯一标识"),
            ("user_id", "UUID", "否", "NULL", "外键→users.id", "关联用户"),
            ("product_id", "INTEGER", "否", "NULL", "外键→products.id", "关联产品"),
            ("alert_type", "VARCHAR(10)", "否", "NULL", "", "预警类型（above/below）"),
            ("threshold", "NUMERIC", "否", "NULL", "", "预警阈值"),
            ("is_active", "BOOLEAN", "否", "true", "", "是否启用"),
            ("created_at", "TIMESTAMPTZ", "否", "now()", "", "创建时间"),
            ("updated_at", "TIMESTAMPTZ", "否", "now()", "", "更新时间"),
        ]
    },
    "alert_logs": {
        "cn": "预警触发记录表",
        "desc": "记录价格预警被触发的历史记录。",
        "columns": [
            ("id", "INTEGER", "否", "自增", "主键", "记录唯一标识"),
            ("alert_id", "INTEGER", "否", "NULL", "外键→price_alerts.id", "关联预警规则"),
            ("triggered_at", "TIMESTAMPTZ", "否", "now()", "", "触发时间"),
            ("price_value", "NUMERIC", "否", "NULL", "", "触发时价格"),
            ("threshold_value", "NUMERIC", "否", "NULL", "", "触发时阈值"),
            ("alert_type", "VARCHAR(10)", "否", "NULL", "", "预警类型"),
            ("product_name", "VARCHAR(255)", "否", "NULL", "", "产品名称（冗余）"),
            ("is_read", "BOOLEAN", "否", "false", "", "是否已读"),
        ]
    },
}

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_font(cell, text, bold=False, size=10, color=None):
    cell.text = text
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0] if para.runs else para.add_run(text)
    run.text = text
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

doc = Document()

# 页面设置
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# 标题
title = doc.add_heading('数据库设计说明', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('本文档描述农产品价格分析系统的数据库表结构设计，包含各表的字段定义、约束关系及说明。')
doc.add_paragraph()

# 表关联关系说明
doc.add_heading('一、表关联关系', level=1)
rel_para = doc.add_paragraph()
rel_para.add_run('系统共包含 9 张业务表，主要关联关系如下：\n').bold = False
relations = [
    'products.category_id  →  categories.id（产品所属分类）',
    'price_records.product_id  →  products.id（价格记录关联产品）',
    'price_records.market_id  →  markets.id（价格记录关联市场）',
    'price_records.source_id  →  data_sources.id（价格记录关联数据来源）',
    'crawl_logs.source_id  →  data_sources.id（爬虫日志关联数据来源）',
    'price_alerts.user_id  →  users.id（预警规则关联用户）',
    'price_alerts.product_id  →  products.id（预警规则关联产品）',
    'alert_logs.alert_id  →  price_alerts.id（预警记录关联预警规则）',
]
for r in relations:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph()
doc.add_heading('二、各表结构详细说明', level=1)

HEADERS = ['字段名', '数据类型', '允许空', '默认值', '约束', '说明']
COL_WIDTHS = [Cm(2.8), Cm(3.2), Cm(1.5), Cm(2.8), Cm(3.2), Cm(4.0)]
HEADER_BG = '4472C4'
HEADER_FG = (255, 255, 255)
ROW_BG_ODD = 'FFFFFF'
ROW_BG_EVEN = 'DCE6F1'

for i, (table_key, info) in enumerate(SCHEMA.items(), 1):
    doc.add_heading(f'{i}. {info["cn"]}（{table_key}）', level=2)
    doc.add_paragraph(info['desc'])

    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置列宽
    for j, width in enumerate(COL_WIDTHS):
        for row in table.rows:
            row.cells[j].width = width

    # 表头
    hdr_row = table.rows[0]
    for j, h in enumerate(HEADERS):
        cell = hdr_row.cells[j]
        set_cell_bg(cell, HEADER_BG)
        set_cell_font(cell, h, bold=True, size=10, color=HEADER_FG)

    # 数据行
    for row_idx, col in enumerate(info['columns']):
        row = table.add_row()
        bg = ROW_BG_ODD if row_idx % 2 == 0 else ROW_BG_EVEN
        for j, val in enumerate(col):
            cell = row.cells[j]
            set_cell_bg(cell, bg)
            is_pk = 'PK' in val or '主键' in val
            set_cell_font(cell, str(val), bold=(j == 0), size=9,
                         color=(0x1F, 0x49, 0x7D) if is_pk else None)

    doc.add_paragraph()

output = r'g:\个人\AgriPriceAnalysis\doc\数据库.docx'
doc.save(output)
print(f'已保存到: {output}')
