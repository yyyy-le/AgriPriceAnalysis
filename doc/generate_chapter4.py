"""生成第四章系统设计文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_run_font(run, name_cn="宋体", name_en="Times New Roman", size=12, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name_cn)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
    return heading


def add_para(doc, text, indent=False, bold=False, font_size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = Pt(22)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if indent:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=font_size, bold=bold)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=True)
    return p


def create_env_table(doc):
    """4.1.1 开发环境配置表"""
    headers = ["序号", "类别", "环境"]
    rows_data = [
        ("1", "操作系统", "Windows 10"),
        ("2", "开发工具", "VS Code / PyCharm"),
        ("3", "数据库", "PostgreSQL 15"),
        ("4", "浏览器", "Microsoft Edge"),
        ("5", "服务器", "Uvicorn 0.34"),
    ]
    table = doc.add_table(rows=1 + len(rows_data), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 设置列宽
    widths = [Cm(2), Cm(4), Cm(6)]
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.width = widths[j]

    # 表头
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=11, bold=True)
        set_cell_shading(cell, "4472C4")
        run.font.color.rgb = RGBColor(255, 255, 255)

    # 数据行
    for i, (no, cat, env) in enumerate(rows_data):
        for j, val in enumerate((no, cat, env)):
            cell = table.cell(i + 1, j)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(val)
            set_run_font(run, size=11)
            if i % 2 == 0:
                set_cell_shading(cell, "DEEAF1")
    return table


def create_arch_diagram(doc):
    """4.1.2 系统实现架构图（分层表格）"""
    layers = [
        ("农产品物价分析系统视图层", "Vue 3", "Element Plus", "4472C4"),
        ("农产品物价分析系统控制层", "FastAPI Router层", None, "5B9BD5"),
        ("农产品物价分析系统逻辑层", "Service层", None, "2E75B6"),
        ("农产品物价分析系统持久层", "SQLAlchemy/ORM层", None, "1F4E79"),
    ]

    for label, tech1, tech2, color in layers:
        table = doc.add_table(rows=1, cols=3 if tech2 else 2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 左侧主标签
        c0 = table.cell(0, 0)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = c0.paragraphs[0].add_run(label)
        set_run_font(run, size=11, bold=True, color=(255, 255, 255))
        set_cell_shading(c0, color)

        # 右侧技术标签
        c1 = table.cell(0, 1)
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = c1.paragraphs[0].add_run(tech1)
        set_run_font(run1, size=11, bold=True)
        set_cell_shading(c1, "D6E4F0")

        if tech2:
            c2 = table.cell(0, 2)
            c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = c2.paragraphs[0].add_run(tech2)
            set_run_font(run2, size=11, bold=True)
            set_cell_shading(c2, "D6E4F0")

        doc.add_paragraph()  # 间距

    # 数据库层
    db_table = doc.add_table(rows=1, cols=1)
    db_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = db_table.cell(0, 0)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = c.paragraphs[0].add_run("数据库服务器 PostgreSQL + Redis")
    set_run_font(run, size=12, bold=True, color=(255, 255, 255))
    set_cell_shading(c, "1F4E79")


def create_db_table(doc, caption_text, col_headers, rows_data):
    """通用数据库表结构表"""
    add_caption(doc, caption_text)
    table = doc.add_table(rows=1 + len(rows_data), cols=len(col_headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for j, h in enumerate(col_headers):
        cell = table.cell(0, j)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=11, bold=True)
        set_cell_shading(cell, "4472C4")
        run.font.color.rgb = RGBColor(255, 255, 255)

    for i, row_vals in enumerate(rows_data):
        for j, val in enumerate(row_vals):
            cell = table.cell(i + 1, j)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10.5)
            if i % 2 == 0:
                set_cell_shading(cell, "DEEAF1")

    doc.add_paragraph()


def create_api_table(doc, caption_text, rows_data):
    """API接口说明表"""
    add_caption(doc, caption_text)
    headers = ["接口名称", "请求方式", "URL路径", "功能说明"]
    table = doc.add_table(rows=1 + len(rows_data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    col_widths = [Cm(3), Cm(2), Cm(5), Cm(5)]
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            cell.width = col_widths[j]

    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=11, bold=True)
        set_cell_shading(cell, "4472C4")
        run.font.color.rgb = RGBColor(255, 255, 255)

    for i, row_vals in enumerate(rows_data):
        for j, val in enumerate(row_vals):
            cell = table.cell(i + 1, j)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10.5)
            if i % 2 == 0:
                set_cell_shading(cell, "DEEAF1")

    doc.add_paragraph()


def main():
    doc = Document()

    # 页面设置：A4，页边距2.5cm
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ========== 第四章标题 ==========
    add_heading_styled(doc, "第四章 系统设计", level=1)

    # ========== 4.1 总体设计 ==========
    add_heading_styled(doc, "4.1 总体设计", level=2)

    # ========== 4.1.1 系统开发环境 ==========
    add_heading_styled(doc, "4.1.1 系统开发环境", level=3)
    add_para(doc, (
        "在正式开发系统前需进行总体设计，确保系统开发环境能保证交易平台正常开发。"
        "本系统的开发环境配置如表4.1所示。"
    ), indent=True)
    doc.add_paragraph()
    add_caption(doc, "表4.1 开发环境配置表")
    create_env_table(doc)
    doc.add_paragraph()

    # ========== 4.1.2 系统实现架构 ==========
    add_heading_styled(doc, "4.1.2 系统实现架构", level=3)
    add_para(doc, (
        "本农产品物价分析系统采用前后端分离架构，后端基于FastAPI框架，前端采用Vue 3框架，"
        "数据库使用PostgreSQL，缓存使用Redis。系统整体遵循分层设计原则，分为视图层、控制层、"
        "逻辑层和持久层，各层职责明确，便于开发和维护，如图4.1所示。"
    ), indent=True)
    doc.add_paragraph()
    create_arch_diagram(doc)
    add_caption(doc, "图4.1 系统实现架构图")
    doc.add_paragraph()
    add_para(doc, (
        "视图层采用Vue 3框架和Element Plus组件库进行界面渲染，构建前端页面。前端页面将数据请求"
        "通过HTTP协议传给控制层，控制层接收到数据之后进行分析，将数据传给逻辑层进行逻辑处理，"
        "然后通过SQLAlchemy ORM持久层去访问数据库，对数据库进行相应操作，最后将逻辑处理的结果"
        "返回给控制层，再由控制层传给前端页面进行展示。"
    ), indent=True)

    # ========== 4.2 功能模块设计 ==========
    add_heading_styled(doc, "4.2 功能模块设计", level=2)
    add_para(doc, (
        "根据第三章的需求分析，系统功能模块分为用户端和管理员端两大部分。用户端包括数据总览、"
        "价格查询、数据分析与预测、AI智能助手、价格预警和个人中心六个模块；管理员端包括用户管理、"
        "数据管理、爬虫管理、数据导入和系统日志五个模块。各模块之间通过统一的RESTful API进行交互。"
    ), indent=True)

    # ========== 4.2.1 用户端功能模块 ==========
    add_heading_styled(doc, "4.2.1 用户端功能模块", level=3)
    add_para(doc, (
        "（1）数据总览模块：展示系统核心统计指标，包括价格记录总量、产品数量、分类数量和市场数量，"
        "同时以折线图展示近30天每日全品类均价走势，以饼图展示各分类数据占比，以柱状图展示今日"
        "最贵/最便宜产品Top10，以雷达图展示价格波动最大的Top8产品，帮助用户快速掌握市场全貌。"
    ), indent=True)
    add_para(doc, (
        "（2）价格查询模块：提供多维度价格数据检索功能，支持按产品名称关键词搜索、按一级/二级分类"
        "筛选，以分页列表形式展示价格记录，包含产品名称、分类、市场、最低价、最高价、均价及"
        "记录时间等字段，满足用户精准查询需求。"
    ), indent=True)
    add_para(doc, (
        "（3）数据分析与预测模块：基于Facebook Prophet时间序列预测算法，对指定农产品的历史价格"
        "数据进行建模，预测未来3至30天的价格走势，并给出预测区间上下界。同时支持调用AI大模型"
        "生成专业的价格分析报告，为采购商和农户提供决策参考。"
    ), indent=True)
    add_para(doc, (
        "（4）AI智能助手模块：集成DeepSeek大语言模型，提供流式对话功能。AI助手内置农产品价格"
        "查询工具，能够理解用户自然语言提问，自动调用数据库查询接口获取实时价格数据，并以"
        "对话形式返回分析结果。"
    ), indent=True)
    add_para(doc, (
        "（5）价格预警模块：允许用户为关注的农产品设置价格阈值预警，支持价格高于阈值（above）"
        "和低于阈值（below）两种预警类型。系统每日自动检测价格数据，触发预警时记录预警日志，"
        "用户可查看预警历史记录并标记已读。"
    ), indent=True)
    add_para(doc, (
        "（6）个人中心模块：提供用户基本信息查看与修改功能，包括用户名、手机号的更新以及密码修改，"
        "保障用户账户安全。"
    ), indent=True)

    # ========== 4.2.2 管理员端功能模块 ==========
    add_heading_styled(doc, "4.2.2 管理员端功能模块", level=3)
    add_para(doc, (
        "（1）用户管理模块：管理员可查看系统全部用户列表，支持关键词搜索，可对用户进行新增、"
        "编辑、启用/禁用状态切换及管理员权限分配等操作，实现对系统用户的全面管控。"
    ), indent=True)
    add_para(doc, (
        "（2）数据管理模块：管理员可浏览所有农产品价格记录，支持按产品聚合展示，可展开查看"
        "某产品的详细价格记录，并支持对单条记录进行编辑和删除操作，保证数据质量。"
    ), indent=True)
    add_para(doc, (
        "（3）爬虫管理模块：管理员可触发新发地农产品价格爬虫任务，支持设置起始页码，并在爬取"
        "过程中实时查看进度（当前页/总页数、已保存/已跳过记录数），支持暂停、恢复和取消操作。"
    ), indent=True)
    add_para(doc, (
        "（4）数据导入模块：支持管理员上传CSV格式的价格数据文件，系统自动解析并批量导入，"
        "对重复数据进行去重处理，返回导入成功和跳过的记录数量统计。"
    ), indent=True)
    add_para(doc, (
        "（5）系统日志模块：按数据来源（爬虫采集、CSV导入等）统计价格记录数量及最近采集时间，"
        "帮助管理员掌握数据采集情况。"
    ), indent=True)

    # ========== 4.3 数据库设计 ==========
    add_heading_styled(doc, "4.3 数据库设计", level=2)
    add_para(doc, (
        "本系统采用PostgreSQL 15作为关系型数据库，使用SQLAlchemy 2.0异步ORM进行数据访问，"
        "通过Alembic进行数据库版本迁移管理。根据系统功能需求，设计了用户表、分类表、市场表、"
        "产品表、价格记录表、价格预警表和预警日志表共七张核心数据表。"
    ), indent=True)

    # 用户表
    add_heading_styled(doc, "4.3.1 用户表（users）", level=3)
    add_para(doc, "用户表存储系统注册用户的基本信息，包括账号、密码、手机号及状态等字段，如表4.2所示。", indent=True)
    create_db_table(doc, "表4.2 用户表（users）结构",
        ["字段名", "类型", "约束", "说明"],
        [
            ("id", "UUID", "PK", "用户唯一标识"),
            ("username", "VARCHAR", "UNIQUE, NOT NULL", "用户名"),
            ("password", "VARCHAR", "NOT NULL", "bcrypt加密密码"),
            ("cellphone", "VARCHAR", "UNIQUE, NOT NULL", "手机号"),
            ("state", "ENUM", "NOT NULL", "账号状态：enabled/disabled"),
            ("is_admin", "BOOLEAN", "DEFAULT FALSE", "是否为管理员"),
            ("created_at", "TIMESTAMP", "DEFAULT NOW()", "创建时间"),
            ("updated_at", "TIMESTAMP", "", "更新时间"),
            ("deleted_at", "TIMESTAMP", "", "软删除时间"),
        ]
    )

    # 分类表
    add_heading_styled(doc, "4.3.2 分类表（categories）", level=3)
    add_para(doc, "分类表采用自关联结构，支持多级分类，通过parent_id字段实现层级关系，如表4.3所示。", indent=True)
    create_db_table(doc, "表4.3 分类表（categories）结构",
        ["字段名", "类型", "约束", "说明"],
        [
            ("id", "INTEGER", "PK, AUTO_INCREMENT", "分类ID"),
            ("name", "VARCHAR", "NOT NULL", "分类名称"),
            ("parent_id", "INTEGER", "FK(categories.id)", "父分类ID，NULL为顶级"),
            ("created_at", "TIMESTAMP", "DEFAULT NOW()", "创建时间"),
        ]
    )

    # 市场表
    add_heading_styled(doc, "4.3.3 市场表（markets）", level=3)
    add_para(doc, "市场表存储农产品交易市场的基本信息，包括市场名称及所在省市，如表4.4所示。", indent=True)
    create_db_table(doc, "表4.4 市场表（markets）结构",
        ["字段名", "类型", "约束", "说明"],
        [
            ("id", "INTEGER", "PK, AUTO_INCREMENT", "市场ID"),
            ("name", "VARCHAR", "NOT NULL", "市场名称"),
            ("province", "VARCHAR", "", "所在省份"),
            ("city", "VARCHAR", "", "所在城市"),
            ("created_at", "TIMESTAMP", "DEFAULT NOW()", "创建时间"),
        ]
    )

    # 产品表
    add_heading_styled(doc, "4.3.4 产品表（products）", level=3)
    add_para(doc, "产品表存储农产品基本信息，通过category_id与分类表关联，如表4.5所示。", indent=True)
    create_db_table(doc, "表4.5 产品表（products）结构",
        ["字段名", "类型", "约束", "说明"],
        [
            ("id", "INTEGER", "PK, AUTO_INCREMENT", "产品ID"),
            ("name", "VARCHAR", "NOT NULL", "产品名称"),
            ("category_id", "INTEGER", "FK(categories.id)", "所属分类"),
            ("unit", "VARCHAR", "DEFAULT '公斤'", "计量单位"),
            ("remark", "TEXT", "", "备注信息"),
            ("created_at", "TIMESTAMP", "DEFAULT NOW()", "创建时间"),
        ]
    )

    # 价格记录表
    add_heading_styled(doc, "4.3.5 价格记录表（price_records）", level=3)
    add_para(doc, (
        "价格记录表是系统的核心数据表，存储农产品在各市场的历史价格数据，"
        "以时间、产品ID和市场ID的组合作为唯一约束，并建立多个索引以提升查询性能，如表4.6所示。"
    ), indent=True)
    create_db_table(doc, "表4.6 价格记录表（price_records）结构",
        ["字段名", "类型", "约束", "说明"],
        [
            ("time", "TIMESTAMP", "PK(联合)", "记录时间"),
            ("product_id", "INTEGER", "PK(联合), FK", "产品ID"),
            ("market_id", "INTEGER", "PK(联合), FK", "市场ID"),
            ("price", "NUMERIC(10,2)", "", "价格"),
            ("min_price", "NUMERIC(10,2)", "", "最低价"),
            ("max_price", "NUMERIC(10,2)", "", "最高价"),
            ("avg_price", "NUMERIC(10,2)", "", "均价"),
            ("source", "VARCHAR", "", "数据来源"),
            ("spec_info", "VARCHAR", "", "规格信息"),
            ("unit_info", "VARCHAR", "", "单位信息"),
        ]
    )

    # 价格预警表
    add_heading_styled(doc, "4.3.6 价格预警表（price_alerts）", level=3)
    add_para(doc, "价格预警表存储用户设置的价格预警规则，支持高于阈值和低于阈值两种预警类型，如表4.7所示。", indent=True)
    create_db_table(doc, "表4.7 价格预警表（price_alerts）结构",
        ["字段名", "类型", "约束", "说明"],
        [
            ("id", "INTEGER", "PK, AUTO_INCREMENT", "预警ID"),
            ("user_id", "UUID", "FK(users.id)", "所属用户"),
            ("product_id", "INTEGER", "FK(products.id)", "监控产品"),
            ("alert_type", "VARCHAR", "NOT NULL", "预警类型：above/below"),
            ("threshold", "NUMERIC", "NOT NULL", "价格阈值"),
            ("is_active", "BOOLEAN", "DEFAULT TRUE", "是否启用"),
            ("created_at", "TIMESTAMP", "DEFAULT NOW()", "创建时间"),
            ("updated_at", "TIMESTAMP", "", "更新时间"),
        ]
    )

    # ========== 4.4 接口设计 ==========
    add_heading_styled(doc, "4.4 接口设计", level=2)
    add_para(doc, (
        "系统后端采用RESTful风格API设计，所有接口统一以/api为前缀，使用JSON格式进行数据交换。"
        "身份认证采用OAuth2 Bearer Token机制，需要鉴权的接口在请求头中携带Authorization: Bearer <token>。"
        "接口按功能模块划分为认证、价格数据、价格预测、价格预警、用户管理和管理员六个路由组。"
    ), indent=True)

    # 认证接口
    add_heading_styled(doc, "4.4.1 认证接口", level=3)
    add_para(doc, "认证模块提供用户登录、退出和Token状态查询功能，如表4.8所示。", indent=True)
    create_api_table(doc, "表4.8 认证接口列表",
        [
            ("用户名密码登录", "POST", "/api/auth/token/password", "提交用户名和密码，返回JWT访问令牌"),
            ("退出登录", "DELETE", "/api/auth/token", "注销当前Token，使其失效"),
            ("查询Token状态", "GET", "/api/auth/token/status", "返回当前Token的有效性及过期时间"),
        ]
    )

    # 价格数据接口
    add_heading_styled(doc, "4.4.2 价格数据接口", level=3)
    add_para(doc, "价格数据模块提供多维度的农产品价格查询与统计功能，如表4.9所示。", indent=True)
    create_api_table(doc, "表4.9 价格数据接口列表",
        [
            ("数据概览统计", "GET", "/api/prices/summary", "返回记录总量、产品数、分类数、市场数及最新更新时间"),
            ("近30天日均价", "GET", "/api/prices/daily-avg", "返回近30天每日全品类均价折线图数据"),
            ("分类数据占比", "GET", "/api/prices/category-stats", "返回各分类价格记录数量及占比"),
            ("今日最贵Top10", "GET", "/api/prices/top-expensive", "返回今日均价最高的10种农产品"),
            ("今日最便宜Top10", "GET", "/api/prices/top-cheapest", "返回今日均价最低的10种农产品"),
            ("价格波动Top8", "GET", "/api/prices/price-volatility", "返回价格波动幅度最大的8种产品"),
            ("各市场数据量", "GET", "/api/prices/market-stats", "返回各市场价格记录数量排行"),
            ("价格列表", "GET", "/api/prices/list", "分页查询价格记录，支持产品名称和分类筛选"),
            ("波动趋势", "GET", "/api/prices/volatility-trend", "返回波动Top8产品近30天每日均价走势"),
        ]
    )

    # 价格预测接口
    add_heading_styled(doc, "4.4.3 价格预测接口", level=3)
    add_para(doc, "价格预测模块基于Prophet算法提供价格预测和AI分析报告功能，如表4.10所示。", indent=True)
    create_api_table(doc, "表4.10 价格预测接口列表",
        [
            ("价格预测数据", "GET", "/api/prices/predict", "基于Prophet模型预测指定产品未来3~30天价格走势"),
            ("AI价格分析报告", "GET", "/api/prices/predict/analysis", "调用DeepSeek大模型生成流式价格分析报告"),
        ]
    )

    # 价格预警接口
    add_heading_styled(doc, "4.4.4 价格预警接口", level=3)
    add_para(doc, "价格预警模块提供预警规则的增删改查及预警日志管理功能，如表4.11所示。", indent=True)
    create_api_table(doc, "表4.11 价格预警接口列表",
        [
            ("获取预警列表", "GET", "/api/alerts/list", "获取当前用户的所有价格预警规则"),
            ("创建预警", "POST", "/api/alerts/create", "为指定产品创建价格阈值预警"),
            ("更新预警", "PUT", "/api/alerts/update/{id}", "修改预警的阈值、类型或启用状态"),
            ("删除预警", "DELETE", "/api/alerts/delete/{id}", "删除指定预警规则"),
            ("预警日志列表", "GET", "/api/alerts/logs", "分页查询当前用户的预警触发记录"),
            ("标记已读", "POST", "/api/alerts/logs/{id}/mark-read", "将指定预警日志标记为已读"),
        ]
    )

    # 管理员接口
    add_heading_styled(doc, "4.4.5 管理员接口", level=3)
    add_para(doc, "管理员模块提供用户管理、数据管理、爬虫控制和数据导入等功能，如表4.12所示。", indent=True)
    create_api_table(doc, "表4.12 管理员接口列表",
        [
            ("用户列表", "GET", "/api/admin/users", "分页查询系统用户，支持关键词搜索"),
            ("新增用户", "POST", "/api/admin/users", "管理员创建新用户账号"),
            ("编辑用户", "PUT", "/api/admin/users/{id}", "修改用户手机号、密码和权限"),
            ("修改用户状态", "PATCH", "/api/admin/users/{id}/state", "启用或禁用用户账号"),
            ("设置管理员", "PATCH", "/api/admin/users/{id}/admin", "授予或撤销管理员权限"),
            ("删除用户", "DELETE", "/api/admin/users/{id}", "删除指定用户"),
            ("产品聚合列表", "GET", "/api/admin/products", "按产品聚合展示价格数据"),
            ("产品价格详情", "GET", "/api/admin/products/{id}/prices", "查看指定产品的详细价格记录"),
            ("编辑价格记录", "PUT", "/api/admin/prices/{id}", "修改单条价格记录"),
            ("删除价格记录", "DELETE", "/api/admin/prices/{id}", "删除单条价格记录"),
            ("触发爬虫", "POST", "/api/crawl/xinfadi", "启动新发地农产品价格爬虫任务"),
            ("暂停爬虫", "POST", "/api/crawl/pause/{task_id}", "暂停正在运行的爬虫任务"),
            ("恢复爬虫", "POST", "/api/crawl/resume/{task_id}", "恢复已暂停的爬虫任务"),
            ("取消爬虫", "POST", "/api/crawl/cancel/{task_id}", "取消爬虫任务"),
            ("查询爬虫状态", "GET", "/api/crawl/status/{task_id}", "获取爬虫任务实时进度"),
            ("CSV数据导入", "POST", "/api/admin/import/csv", "上传CSV文件批量导入价格数据"),
            ("系统日志", "GET", "/api/admin/logs", "按数据来源统计采集记录数量"),
        ]
    )

    # ========== 4.5 安全设计 ==========
    add_heading_styled(doc, "4.5 安全设计", level=2)
    add_para(doc, (
        "系统在安全设计方面采取了多层防护措施。在身份认证方面，采用OAuth2 Bearer Token机制，"
        "使用JWT（JSON Web Token）进行无状态身份验证，Token存储于Redis中并支持主动注销，"
        "有效防止Token泄露后的持续滥用。"
    ), indent=True)
    add_para(doc, (
        "在密码安全方面，用户密码采用bcrypt算法进行哈希加密存储，bcrypt内置盐值机制，"
        "能有效抵御彩虹表攻击，即使数据库泄露也无法直接还原明文密码。"
    ), indent=True)
    add_para(doc, (
        "在接口安全方面，系统引入fastapi-limiter对接口进行访问频率限制，防止暴力破解和"
        "恶意刷接口攻击。管理员接口通过is_admin字段进行权限校验，普通用户无法访问管理功能。"
    ), indent=True)
    add_para(doc, (
        "在数据安全方面，所有数据库操作均通过SQLAlchemy参数化查询执行，从根本上杜绝SQL注入风险。"
        "系统对用户输入进行严格的Pydantic模型校验，确保数据格式合法性。"
    ), indent=True)

    # ========== 4.6 本章小结 ==========
    add_heading_styled(doc, "4.6 本章小结", level=2)
    add_para(doc, (
        "本章对农产品物价数据分析系统进行了全面的系统设计。首先介绍了系统的开发环境配置，"
        "明确了操作系统、开发工具、数据库和服务器等基础环境；然后阐述了基于前后端分离的"
        "系统实现架构，采用Vue 3+FastAPI+PostgreSQL+Redis的技术栈，遵循视图层、控制层、"
        "逻辑层和持久层的分层设计原则。在功能模块设计部分，详细描述了用户端六大模块和"
        "管理员端五大模块的设计方案；在数据库设计部分，给出了七张核心数据表的结构设计；"
        "在接口设计部分，按模块整理了系统全部RESTful API接口；最后从身份认证、密码安全、"
        "接口防护和数据安全四个维度阐述了系统的安全设计策略。以上设计为后续系统实现提供了"
        "完整的技术方案。"
    ), indent=True)

    output_path = r"g:\个人\AgriPriceAnalysis\doc\第四章系统设计.docx"
    doc.save(output_path)
    print(f"文档已保存至: {output_path}")


if __name__ == "__main__":
    main()
