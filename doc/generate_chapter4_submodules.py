# -*- coding: utf-8 -*-
"""生成第四章子模块设计文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_run_font(run, name_cn="宋体", name_en="Times New Roman", size=12, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name_cn)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
    return heading


def add_para(doc, text, indent=False, bold=False, font_size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = Pt(22)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if indent:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=font_size, bold=bold)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=True)
    return p


def add_module_table(doc, caption, rows_data):
    add_caption(doc, caption)
    table = doc.add_table(rows=1 + len(rows_data), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    col_widths = [Cm(1.5), Cm(3), Cm(10)]
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            cell.width = col_widths[j]

    headers = ["序号", "模块名称", "用户操作"]
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=11, bold=True)
        set_cell_shading(cell, "4472C4")
        run.font.color.rgb = RGBColor(255, 255, 255)

    for i, (no, name, content) in enumerate(rows_data):
        c0 = table.cell(i + 1, 0)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = c0.paragraphs[0].add_run(str(no))
        set_run_font(run, size=11)

        c1 = table.cell(i + 1, 1)
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = c1.paragraphs[0].add_run(name)
        set_run_font(run, size=11)

        c2 = table.cell(i + 1, 2)
        lines = content.split('\n')
        for k, line in enumerate(lines):
            if k == 0:
                p = c2.paragraphs[0]
            else:
                p = c2.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            run = p.add_run(line)
            set_run_font(run, size=11)

        if i % 2 == 0:
            set_cell_shading(c0, "DEEAF1")
            set_cell_shading(c1, "DEEAF1")
            set_cell_shading(c2, "DEEAF1")

    doc.add_paragraph()


def main():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    add_heading_styled(doc, "4.2 子模块详细设计", level=2)

    # ===== 4.2.1 用户登录与注册模块 =====
    add_heading_styled(doc, "4.2.1 用户登录与注册模块", level=3)
    add_para(doc, "用户登录与注册模块需要简单易懂，易于操作。具体功能如下：", indent=True)
    doc.add_paragraph()

    add_para(doc, "（1）用户登录详细设计如下表4.2所示：", indent=False)
    add_module_table(doc, "表4.2 用户登录详细功能设计表", [
        (1, "功能描述", "用户进行登录，后台验证通过后跳转首页"),
        (2, "输入项", "输入\n1.用户名 string\n2.密码 string"),
        (3, "处理描述",
            "1.用户提交信息\n"
            "2.系统校验用户名和密码是否为空\n"
            "3.调用后端 /api/auth/token/password 接口进行身份验证\n"
            "4.验证通过后将JWT Token存入本地，跳转至首页"),
        (4, "处理结果", "如果验证通过，则直接跳转至首页正文页面。"),
        (5, "异常处理", "通过弹窗提示用户错误信息。"),
    ])

    add_para(doc, "（2）用户注册详细设计如下表4.3所示：", indent=False)
    add_module_table(doc, "表4.3 用户注册详细功能设计表", [
        (1, "功能描述", "用户进行账号注册，注册成功后可登录系统"),
        (2, "输入项", "输入\n1.用户名 string\n2.手机号 string\n3.密码 string\n4.确认密码 string"),
        (3, "处理描述",
            "1.用户填写注册信息并提交\n"
            "2.系统校验用户名、手机号、密码是否为空\n"
            "3.系统校验两次密码是否一致\n"
            "4.调用注册接口，后台检查用户名和手机号是否已存在\n"
            "5.满足条件后创建用户账号，密码经bcrypt加密存储"),
        (4, "处理结果", "注册成功后提示用户，并跳转至登录页面。"),
        (5, "异常处理", "通过弹窗提示用户名已存在、手机号已注册等错误信息。"),
    ])

    # ===== 4.2.2 数据总览模块 =====
    add_heading_styled(doc, "4.2.2 数据总览模块", level=3)
    add_para(doc, "数据总览模块为用户提供农产品价格市场的全局视图，登录后默认展示该页面。具体功能如下：", indent=True)
    doc.add_paragraph()

    add_para(doc, "（1）数据概览统计详细设计如下表4.4所示：", indent=False)
    add_module_table(doc, "表4.4 数据概览统计详细功能设计表", [
        (1, "功能描述", "展示系统核心统计指标，包括价格记录总量、产品数、分类数、市场数及最新更新时间"),
        (2, "输入项", "无（用户登录后自动加载）"),
        (3, "处理描述",
            "1.页面加载时调用 GET /api/prices/summary 接口\n"
            "2.后台查询 price_records、products、categories、markets 表的记录总数\n"
            "3.返回统计数据并渲染至卡片组件"),
        (4, "处理结果", "页面展示总记录数、产品数、分类数、市场数四个统计卡片及最新数据更新时间。"),
        (5, "异常处理", "接口请求失败时显示加载错误提示，数据为空时显示0。"),
    ])

    add_para(doc, "（2）近30天日均价走势详细设计如下表4.5所示：", indent=False)
    add_module_table(doc, "表4.5 近30天日均价走势详细功能设计表", [
        (1, "功能描述", "以折线图展示近30天每日全品类均价走势"),
        (2, "输入项", "无（自动加载）"),
        (3, "处理描述",
            "1.调用 GET /api/prices/daily-avg 接口\n"
            "2.后台按日期分组计算所有品类均价，取近30天数据\n"
            "3.前端使用ECharts渲染折线图，X轴为日期，Y轴为均价"),
        (4, "处理结果", "展示近30天价格走势折线图，支持鼠标悬停查看具体日期和价格。"),
        (5, "异常处理", "数据不足时图表显示空状态提示。"),
    ])

    add_para(doc, "（3）今日价格排行详细设计如下表4.6所示：", indent=False)
    add_module_table(doc, "表4.6 今日价格排行详细功能设计表", [
        (1, "功能描述", "展示今日均价最贵和最便宜的农产品各Top10"),
        (2, "输入项", "无（自动加载）"),
        (3, "处理描述",
            "1.分别调用 GET /api/prices/top-expensive 和 GET /api/prices/top-cheapest\n"
            "2.后台查询当日最新价格记录，按均价降序/升序排列取前10条\n"
            "3.前端以柱状图形式展示"),
        (4, "处理结果", "展示今日最贵Top10和最便宜Top10产品的均价柱状图。"),
        (5, "异常处理", "当日无数据时提示「暂无今日价格数据」。"),
    ])

    # ===== 4.2.3 价格查询模块 =====
    add_heading_styled(doc, "4.2.3 价格查询模块", level=3)
    add_para(doc, "价格查询模块提供多维度的农产品价格数据检索功能。具体功能如下：", indent=True)
    doc.add_paragraph()

    add_para(doc, "（1）价格列表查询详细设计如下表4.7所示：", indent=False)
    add_module_table(doc, "表4.7 价格列表查询详细功能设计表", [
        (1, "功能描述", "用户通过分类筛选或关键词搜索查询农产品价格记录"),
        (2, "输入项",
            "输入\n"
            "1.产品名称关键词 string（可选）\n"
            "2.一级分类 ID（可选）\n"
            "3.二级分类 ID（可选）\n"
            "4.页码 int，每页条数 int"),
        (3, "处理描述",
            "1.用户输入搜索条件后点击查询\n"
            "2.调用 GET /api/prices/list 接口，传入筛选参数\n"
            "3.后台动态拼接SQL，支持产品名称模糊匹配和分类过滤\n"
            "4.返回分页数据，包含产品名、分类、市场、最低价、最高价、均价、时间"),
        (4, "处理结果", "以分页表格展示价格记录，支持翻页浏览。"),
        (5, "异常处理", "无匹配结果时显示「暂无数据」，接口异常时弹窗提示。"),
    ])

    # ===== 4.2.4 数据分析与预测模块 =====
    add_heading_styled(doc, "4.2.4 数据分析与预测模块", level=3)
    add_para(doc, "数据分析与预测模块基于Prophet时间序列算法对农产品价格进行预测分析。具体功能如下：", indent=True)
    doc.add_paragraph()

    add_para(doc, "（1）价格预测详细设计如下表4.8所示：", indent=False)
    add_module_table(doc, "表4.8 价格预测详细功能设计表", [
        (1, "功能描述", "用户选择产品和预测天数，系统基于历史数据预测未来价格走势"),
        (2, "输入项",
            "输入\n"
            "1.产品ID int\n"
            "2.预测天数 int（3~30天）"),
        (3, "处理描述",
            "1.用户选择产品和预测天数后点击预测\n"
            "2.调用 GET /api/prices/predict 接口\n"
            "3.后台查询该产品历史价格数据（至少7天），使用Prophet模型拟合\n"
            "4.生成未来N天逐日预测价格及置信区间上下界\n"
            "5.返回历史数据（近30天）和预测数据，前端以折线图展示"),
        (4, "处理结果", "展示历史价格折线与预测价格折线（含置信区间），并显示预测涨跌幅百分比。"),
        (5, "异常处理", "历史数据不足7天时提示数据不足，无法预测。"),
    ])

    add_para(doc, "（2）AI价格分析报告详细设计如下表4.9所示：", indent=False)
    add_module_table(doc, "表4.9 AI价格分析报告详细功能设计表", [
        (1, "功能描述", "调用DeepSeek大语言模型，基于预测数据生成专业价格分析报告"),
        (2, "输入项",
            "输入\n"
            "1.产品ID int\n"
            "2.预测天数 int"),
        (3, "处理描述",
            "1.用户点击「生成AI分析报告」按钮\n"
            "2.调用 GET /api/prices/predict/analysis 接口（流式响应）\n"
            "3.后台运行Prophet预测，将历史统计和预测数据组装为Prompt\n"
            "4.调用DeepSeek API进行流式生成，逐字返回分析内容\n"
            "5.前端以打字机效果实时展示报告内容"),
        (4, "处理结果", "展示包含走势分析、趋势预判、风险提示和操作建议的完整分析报告。"),
        (5, "异常处理", "AI接口超时或异常时提示「报告生成失败，请稍后重试」。"),
    ])

    # ===== 4.2.5 AI智能助手模块 =====
    add_heading_styled(doc, "4.2.5 AI智能助手模块", level=3)
    add_para(doc, "AI智能助手模块集成DeepSeek大语言模型，提供自然语言交互式价格查询功能。具体功能如下：", indent=True)
    doc.add_paragraph()

    add_para(doc, "（1）智能对话详细设计如下表4.10所示：", indent=False)
    add_module_table(doc, "表4.10 AI智能对话详细功能设计表", [
        (1, "功能描述", "用户以自然语言提问，AI助手自动调用数据库工具查询价格并回答"),
        (2, "输入项",
            "输入\n"
            "1.用户消息 string\n"
            "2.历史对话记录 list（上下文）"),
        (3, "处理描述",
            "1.用户输入问题后发送\n"
            "2.调用 POST /api/ai/chat 接口（流式响应）\n"
            "3.后台将用户消息和历史上下文发送给DeepSeek模型\n"
            "4.模型判断是否需要调用工具（如查询某产品价格、获取分类列表等）\n"
            "5.若需要工具调用，后台执行数据库查询并将结果返回给模型\n"
            "6.模型生成最终回答，以流式方式逐字返回前端"),
        (4, "处理结果", "以对话气泡形式展示AI回答，支持Markdown格式渲染。"),
        (5, "异常处理", "网络异常或模型超时时提示「AI助手暂时不可用」。"),
    ])

    # ===== 4.2.6 价格预警模块 =====
    add_heading_styled(doc, "4.2.6 价格预警模块", level=3)
    add_para(doc, "价格预警模块允许用户为关注的农产品设置价格阈值，系统自动监控并记录预警触发情况。具体功能如下：", indent=True)
    doc.add_paragraph()

    add_para(doc, "（1）创建价格预警详细设计如下表4.11所示：", indent=False)
    add_module_table(doc, "表4.11 创建价格预警详细功能设计表", [
        (1, "功能描述", "用户为指定农产品设置价格阈值预警规则"),
        (2, "输入项",
            "输入\n"
            "1.产品ID int\n"
            "2.预警类型 string（above-高于阈值/below-低于阈值）\n"
            "3.价格阈值 float"),
        (3, "处理描述",
            "1.用户选择产品、预警类型并输入阈值后提交\n"
            "2.调用 POST /api/alerts/create 接口\n"
            "3.后台验证产品存在性，将预警规则写入 price_alerts 表\n"
            "4.立即检查今日价格是否已触发预警，若触发则写入 alert_logs 表"),
        (4, "处理结果", "预警创建成功，若今日价格已触发则同时提示「今日价格已触发预警」。"),
        (5, "异常处理", "产品不存在或参数非法时返回错误提示。"),
    ])

    add_para(doc, "（2）预警日志查询详细设计如下表4.12所示：", indent=False)
    add_module_table(doc, "表4.12 预警日志查询详细功能设计表", [
        (1, "功能描述", "用户查看历史价格预警触发记录，并可标记已读"),
        (2, "输入项",
            "输入\n"
            "1.页码 int\n"
            "2.每页条数 int"),
        (3, "处理描述",
            "1.页面加载时调用 GET /api/alerts/logs 接口\n"
            "2.后台关联 alert_logs 和 price_alerts 表，过滤当前用户的记录\n"
            "3.返回触发时间、产品名称、触发价格、阈值、预警类型和已读状态\n"
            "4.用户点击标记已读时调用 POST /api/alerts/logs/{id}/mark-read"),
        (4, "处理结果", "以分页列表展示预警记录，未读记录高亮显示，标记已读后状态更新。"),
        (5, "异常处理", "无预警记录时显示「暂无预警记录」。"),
    ])

    # ===== 4.2.7 个人中心模块 =====
    add_heading_styled(doc, "4.2.7 个人中心模块", level=3)
    add_para(doc, "个人中心模块提供用户基本信息管理功能。具体功能如下：", indent=True)
    doc.add_paragraph()

    add_para(doc, "（1）用户个人信息编辑功能设计如下表4.13所示：", indent=False)
    add_module_table(doc, "表4.13 用户个人信息编辑功能设计表", [
        (1, "功能描述", "用户进行个人信息修改，权限验证通过后完成修改"),
        (2, "输入项",
            "输入\n"
            "1.用户名 string\n"
            "2.手机号 string"),
        (3, "处理描述",
            "1.用户输入修改信息并提交\n"
            "2.系统判断用户名、手机号等信息是否为空\n"
            "3.系统验证输入的手机号是否已被其他用户占用\n"
            "4.满足条件后进行信息修改\n"
            "5.系统更新缓存用户信息"),
        (4, "处理结果", "如果更新通过，则直接显示已更新数据的个人信息页面。"),
        (5, "异常处理", "通过弹窗提示用户错误信息。"),
    ])

    add_para(doc, "（2）修改密码功能设计如下表4.14所示：", indent=False)
    add_module_table(doc, "表4.14 修改密码功能设计表", [
        (1, "功能描述", "用户修改登录密码，需验证原密码后方可修改"),
        (2, "输入项",
            "输入\n"
            "1.原密码 string\n"
            "2.新密码 string\n"
            "3.确认新密码 string"),
        (3, "处理描述",
            "1.用户输入原密码和新密码后提交\n"
            "2.系统校验新密码与确认密码是否一致\n"
            "3.调用接口验证原密码是否正确\n"
            "4.原密码验证通过后，使用bcrypt对新密码加密并更新"),
        (4, "处理结果", "密码修改成功后提示用户，并要求重新登录。"),
        (5, "异常处理", "原密码错误或两次新密码不一致时弹窗提示。"),
    ])

    # ===== 4.2.8 用户管理模块（管理员）=====
    add_heading_styled(doc, "4.2.8 用户管理模块（管理员）", level=3)
    add_para(doc, "用户管理模块供管理员对系统用户进行全面管控。具体功能如下：", indent=True)
    doc.add_paragraph()

    add_para(doc, "（1）用户列表与管理详细设计如下表4.15所示：", indent=False)
    add_module_table(doc, "表4.15 用户列表与管理详细功能设计表", [
        (1, "功能描述", "管理员查看系统全部用户，并进行新增、编辑、状态管理等操作"),
        (2, "输入项",
            "输入\n"
            "1.搜索关键词 string（可选，匹配用户名或手机号）\n"
            "2.页码 int，每页条数 int"),
        (3, "处理描述",
            "1.调用 GET /api/admin/users 接口获取用户列表\n"
            "2.管理员可点击新增用户，填写用户名、手机号、密码和权限后提交\n"
            "3.管理员可对用户进行启用/禁用操作，调用 PATCH /api/admin/users/{id}/state\n"
            "4.管理员可授予或撤销管理员权限，调用 PATCH /api/admin/users/{id}/admin\n"
            "5.管理员可删除用户，调用 DELETE /api/admin/users/{id}"),
        (4, "处理结果", "用户列表实时刷新，操作结果以提示框反馈。"),
        (5, "异常处理", "用户名或手机号重复时提示具体冲突信息。"),
    ])

    # ===== 4.2.9 爬虫管理模块（管理员）=====
    add_heading_styled(doc, "4.2.9 爬虫管理模块（管理员）", level=3)
    add_para(doc, "爬虫管理模块供管理员控制新发地农产品价格数据的自动采集任务。具体功能如下：", indent=True)
    doc.add_paragraph()

    add_para(doc, "（1）爬虫任务控制详细设计如下表4.16所示：", indent=False)
    add_module_table(doc, "表4.16 爬虫任务控制详细功能设计表", [
        (1, "功能描述", "管理员启动、暂停、恢复或取消新发地价格数据爬取任务"),
        (2, "输入项",
            "输入\n"
            "1.起始页码 int（默认为1）"),
        (3, "处理描述",
            "1.管理员设置起始页码后点击启动，调用 POST /api/crawl/xinfadi\n"
            "2.后台生成任务ID，以后台异步任务方式运行爬虫\n"
            "3.前端轮询 GET /api/crawl/status/{task_id} 获取实时进度\n"
            "4.管理员可点击暂停（POST /api/crawl/pause/{task_id}）或恢复（POST /api/crawl/resume/{task_id}）\n"
            "5.管理员可点击取消终止任务（POST /api/crawl/cancel/{task_id}）"),
        (4, "处理结果", "实时展示当前页/总页数、已保存/已跳过记录数，任务完成后显示汇总结果。"),
        (5, "异常处理", "爬虫异常时任务状态变为failed，展示错误信息。"),
    ])

    # ===== 4.2.10 数据导入模块（管理员）=====
    add_heading_styled(doc, "4.2.10 数据导入模块（管理员）", level=3)
    add_para(doc, "数据导入模块支持管理员通过CSV文件批量导入农产品价格数据。具体功能如下：", indent=True)
    doc.add_paragraph()

    add_para(doc, "（1）CSV数据导入详细设计如下表4.17所示：", indent=False)
    add_module_table(doc, "表4.17 CSV数据导入详细功能设计表", [
        (1, "功能描述", "管理员上传CSV格式价格数据文件，系统解析并批量写入数据库"),
        (2, "输入项",
            "输入\n"
            "1.CSV文件（包含产品名称、分类、市场、均价、最低价、最高价、时间等字段）"),
        (3, "处理描述",
            "1.管理员选择CSV文件后点击上传，调用 POST /api/admin/import/csv\n"
            "2.后台读取CSV内容，逐行解析数据\n"
            "3.自动匹配或创建对应的产品、分类、市场记录\n"
            "4.使用 ON CONFLICT DO NOTHING 对重复数据进行去重处理\n"
            "5.返回成功导入数量和跳过数量统计"),
        (4, "处理结果", "显示导入结果：成功X条，跳过Y条，并列出异常行信息。"),
        (5, "异常处理", "文件格式错误或字段缺失时提示具体错误行号和原因。"),
    ])

    output_path = r"g:\个人\AgriPriceAnalysis\doc\第四章子模块设计.docx"
    doc.save(output_path)
    print(f"文档已保存至: {output_path}")


if __name__ == "__main__":
    main()
